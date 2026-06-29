#input_type_name: BuildDocumentRequest
#output_type_name: BuildDocumentResponse
#function_name: build_document
#python_packages: python-docx, docxtpl

import os
import tempfile
from typing import Dict, Any
from pydantic import BaseModel
from lemma_sdk import FunctionContext, Pod

class BuildDocumentRequest(BaseModel):
    entities: Dict[str, Any]
    legal_framework: str
    session_id: str
    document_type: str

class BuildDocumentResponse(BaseModel):
    docx_url: str
    status: str

def create_mock_template(path: str, doc_type: str):
    """Creates a basic Word document with Jinja2 placeholder tags if it does not exist."""
    from docx import Document
    doc = Document()
    doc.add_heading("IN THE JURISDICTIONAL COURT OF INDIA", level=0)
    
    if doc_type == "civil_plaint":
        doc.add_paragraph("CIVIL PLAINT NO. _____ OF 2026")
        doc.add_paragraph("In the matter of:")
        doc.add_paragraph("PLAINTIFF NAME: {{ plaintiff_name }}")
        doc.add_paragraph("PLAINTIFF ADDRESS: {{ plaintiff_address }}")
        doc.add_paragraph("versus")
        doc.add_paragraph("DEFENDANT NAME: {{ defendant_name }}")
        doc.add_paragraph("VALUATION OF SUIT: {{ valuation_of_suit }}")
        doc.add_paragraph("CAUSE OF ACTION DETAILS: {{ cause_of_action_desc }}")
        doc.add_paragraph("PRAYER CLAUSE: {{ relief_sought }}")
    elif doc_type == "bail_application":
        doc.add_paragraph("CRIMINAL BAIL APPLICATION NO. _____ OF 2026")
        doc.add_paragraph("In the matter of:")
        doc.add_paragraph("APPLICANT NAME: {{ plaintiff_name }}")
        doc.add_paragraph("SECTIONS CHARGED: {{ statutory_sections_mentioned }}")
        doc.add_paragraph("COURT: {{ court_name }}")
        doc.add_paragraph("GROUNDS: {{ cause_of_action_desc }}")
    else:
        doc.add_paragraph("LEGAL DEMAND NOTICE")
        doc.add_paragraph("To accused: {{ defendant_name }} residing at {{ defendant_address }}")
        doc.add_paragraph("From complainant: {{ plaintiff_name }}")
        doc.add_paragraph("dishonour_date: {{ cheque_dishonour_date }}")
        
    doc.save(path)

async def build_document(ctx: FunctionContext, data: BuildDocumentRequest) -> BuildDocumentResponse:
    from docxtpl import DocxTemplate
    pod = Pod.from_env()
    
    # 1. Map fields to template parameters
    entities = data.entities
    plaintiff = entities.get("plaintiff", {})
    defendant = entities.get("defendant", {})
    court = entities.get("court", {})
    coa = entities.get("cause_of_action", {})
    
    context = {
        "plaintiff_name": plaintiff.get("name") or "",
        "plaintiff_address": plaintiff.get("address") or "",
        "defendant_name": defendant.get("name") or "",
        "defendant_address": defendant.get("address") or "",
        "court_name": court.get("name") or "",
        "court_district": court.get("district") or "",
        "court_state": court.get("state") or "",
        "cause_of_action_desc": coa.get("description") or "",
        "cause_of_action_date": coa.get("date") or "",
        "relief_sought": ", ".join(entities.get("relief_sought", [])),
        "statutory_sections_mentioned": ", ".join(entities.get("statutory_sections_mentioned", [])),
        "valuation_of_suit": entities.get("valuation_of_suit") or "",
        "cheque_dishonour_date": entities.get("cheque_dishonour_date") or ""
    }

    # Create temporary directory for compiling
    with tempfile.TemporaryDirectory() as tmpdir:
        template_path = os.path.join(tmpdir, f"template_{data.document_type}.docx")
        output_path = os.path.join(tmpdir, f"vaaddoc_{data.session_id}.docx")
        
        # 2. Load template from Lemma files store if set, else write basic mock
        template_loaded = False
        if pod:
            try:
                content = pod.files.download(f"/templates/{data.document_type}.docx")
                with open(template_path, "wb") as f:
                    f.write(content)
                template_loaded = True
            except Exception:
                pass
                
        if not template_loaded:
            create_mock_template(template_path, data.document_type)

        # 3. Render placeholders via docxtpl
        doc = DocxTemplate(template_path)
        doc.render(context)
        doc.save(output_path)

        # 4. Upload generated docx back to Lemma files store
        url = f"/local/{output_path}"
        if pod:
            try:
                ref = pod.files.upload(
                    local_path=output_path,
                    name=f"vaaddoc_{data.session_id}.docx",
                    directory_path="/generated"
                )
                url_resp = pod.files.get_url(ref.path)
                url = url_resp.url
            except Exception:
                # Fallback to local file copying
                fallback_dir = os.path.join(tempfile.gettempdir(), "vaaddoc_generated")
                os.makedirs(fallback_dir, exist_ok=True)
                local_fallback = os.path.join(fallback_dir, f"vaaddoc_{data.session_id}.docx")
                with open(output_path, "rb") as sf, open(local_fallback, "wb") as df:
                    df.write(sf.read())
                url = f"/local/{local_fallback}"

    return BuildDocumentResponse(
        docx_url=url,
        status="complete"
    )
